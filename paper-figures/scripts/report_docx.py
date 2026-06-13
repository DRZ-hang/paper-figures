#!/usr/bin/env python3
"""report_docx.py — assemble the stage-7 figure report as a Word (.docx) document.

Builds the hand-off report: a title, manuscript metadata, an asset index, then one
block per figure/table with the embedded figure image (or inline three-line table),
caption, annotations, in-text citation location, and a reproducibility note.

Output language is selectable via `lang`:
  - "en"        — English only
  - "zh"        — Chinese only / 全中文
  - "bilingual" — both (English + 中文 captions; default)

    from report_docx import build_report
    build_report("Figure_Report.docx", title=..., meta=..., items=..., lang="en")

Each item is a dict. Text fields accept either a plain string or a
``{"en": "...", "zh": "..."}`` dict; when a string is given it is used for every
language (handy when the text is language-neutral). Legacy ``caption_en`` /
``caption_zh`` keys are still accepted.

  figure: {"kind":"figure", "number":"1", "title": {"en":..,"zh":..},
           "image":"figures/Fig1.png", "width_in":6.0,
           "claim": {"en":..,"zh":..}, "caption": {"en":..,"zh":..},
           "annotations": {"en":[...], "zh":[...]},   # or a shared list
           "citation": {"en":..,"zh":..}, "repro":"scripts/make_fig1.py"}
  table:  {"kind":"table", "number":"1", "title": {...}, "csv":"figures/Table1.csv",
           "table_title": {...}, "table_note": {...}, ...same text fields... }

Requires python-docx (and pandas for table CSVs).
"""
from __future__ import annotations

from pathlib import Path

from docx_tables import add_three_line_table

# section labels per language
_LABELS = {
    "bilingual": {"asset": "Asset index / 图表清单", "cols": ["#", "File 文件", "Claim 结论"],
                  "claim": "Claim / 论证: ", "cap_en": "Caption (EN): ", "cap_zh": "图注 (中文): ",
                  "ann": "Annotations / 标注", "cite": "Citation location / 引用位置: ",
                  "repro": "Reproducibility / 可复现: ", "fig": "Figure ", "tbl": "Table "},
    "en": {"asset": "Asset index", "cols": ["#", "File", "Claim"],
           "claim": "Claim: ", "cap": "Caption: ", "ann": "Annotations",
           "cite": "Citation location: ", "repro": "Reproducibility: ",
           "fig": "Figure ", "tbl": "Table "},
    "zh": {"asset": "图表清单", "cols": ["#", "文件", "结论"],
           "claim": "论证:", "cap": "图注:", "ann": "标注",
           "cite": "引用位置:", "repro": "可复现:", "fig": "图 ", "tbl": "表 "},
}


def _pick(val, lang):
    """Resolve a string-or-{en,zh} field to one string for `lang`."""
    if isinstance(val, dict):
        if lang == "zh":
            return val.get("zh") or val.get("en") or ""
        return val.get("en") or val.get("zh") or ""
    return val or ""


def _caption_dict(it):
    """Normalize caption to {en, zh}, supporting legacy caption_en/caption_zh."""
    cap = it.get("caption")
    if isinstance(cap, dict):
        return cap
    if isinstance(cap, str):
        return {"en": cap, "zh": cap}
    return {"en": it.get("caption_en", ""), "zh": it.get("caption_zh", "")}


def _heading(doc, text, size, bold=True, space_before=10, space_after=4, color=None):
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def _labelled(doc, label, text, size=10, label_color=(0x1F, 0x4E, 0x79)):
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(*label_color)
    if text:
        t = p.add_run(text)
        t.font.size = Pt(size)
    return p


def build_report(path, title, subtitle=None, meta=None, items=None,
                 base_dir=None, lang="bilingual", body_font="Calibri", body_pt=10):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if lang not in _LABELS:
        raise ValueError(f"lang must be one of {list(_LABELS)}, got {lang!r}")
    L = _LABELS[lang]
    clang = "zh" if lang == "zh" else "en"  # language for non-caption metadata

    base = Path(base_dir) if base_dir else Path(".")
    doc = Document()
    s = doc.styles["Normal"]
    s.font.name = body_font
    s.font.size = Pt(body_pt)

    _heading(doc, _pick(title, clang), 18, color=(0x1F, 0x4E, 0x79), space_before=0, space_after=2)
    if subtitle:
        _heading(doc, _pick(subtitle, clang), 12, bold=False, space_before=0, space_after=8)

    if meta:
        for label, value in meta:
            _labelled(doc, f"{_pick(label, clang)}: ", _pick(value, clang), size=body_pt)

    if items:
        _heading(doc, L["asset"], 12, space_before=12)
        idx = doc.add_table(rows=1, cols=3)
        idx.style = "Light Grid Accent 1"
        for j, h in enumerate(L["cols"]):
            r = idx.rows[0].cells[j].paragraphs[0].add_run(h)
            r.bold = True
            r.font.size = Pt(body_pt - 1)
        for it in items:
            prefix = L["fig"] if it["kind"] == "figure" else L["tbl"]
            cells = idx.add_row().cells
            fname = Path(it.get("image") or it.get("csv", "")).name
            for j, v in enumerate([f"{prefix}{it['number']}", fname, _pick(it.get('claim'), clang)]):
                rr = cells[j].paragraphs[0].add_run(v)
                rr.font.size = Pt(body_pt - 1)

    for it in items or []:
        prefix = L["fig"] if it["kind"] == "figure" else L["tbl"]
        head = f"{prefix}{it['number']}"
        if it.get("title"):
            head += f" — {_pick(it['title'], clang)}"
        _heading(doc, head, 13, color=(0x1F, 0x4E, 0x79), space_before=14)

        if it["kind"] == "figure" and it.get("image"):
            img = base / it["image"]
            if img.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(img), width=Inches(it.get("width_in", 6.0)))
        elif it["kind"] == "table" and it.get("csv"):
            import pandas as pd
            df = pd.read_csv(base / it["csv"])
            add_three_line_table(doc, df, title=_pick(it.get("table_title"), clang),
                                 note=_pick(it.get("table_note"), clang))

        if it.get("claim"):
            _labelled(doc, L["claim"], _pick(it["claim"], clang))

        cap = _caption_dict(it)
        if lang == "bilingual":
            if cap.get("en"):
                _labelled(doc, L["cap_en"], cap["en"])
            if cap.get("zh"):
                _labelled(doc, L["cap_zh"], cap["zh"])
        else:
            _labelled(doc, L["cap"], cap.get(clang) or cap.get("en") or cap.get("zh"))

        anns = it.get("annotations")
        if isinstance(anns, dict):
            anns = anns.get(clang) or anns.get("en") or anns.get("zh")
        if anns:
            _labelled(doc, L["ann"], "")
            for a in anns:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(1)
                if isinstance(a, (tuple, list)):
                    r = p.add_run(f"{a[0]}: "); r.bold = True; r.font.size = Pt(body_pt)
                    p.add_run(a[1]).font.size = Pt(body_pt)
                else:
                    p.add_run(a).font.size = Pt(body_pt)

        if it.get("citation"):
            _labelled(doc, L["cite"], _pick(it["citation"], clang))
        if it.get("repro"):
            _labelled(doc, L["repro"], _pick(it["repro"], clang))

    doc.save(str(path))
    return path
