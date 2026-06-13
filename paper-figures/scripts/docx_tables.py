#!/usr/bin/env python3
"""docx_tables.py — write a DataFrame as an academic three-line (三线表) Word table.

A three-line table has exactly three rules — a thick rule above the header, a thin
rule under the header, and a thick rule at the bottom — and no vertical lines. This
is the standard table format for most journals and for Chinese theses (三线表).

    from docx_tables import three_line_table
    three_line_table(df, "Table1.docx", title="Table 1. ...", note="n = ...")

Requires python-docx (`pip install python-docx`).
"""
from __future__ import annotations


def _set_cell_border(cell, **edges):
    """Set per-cell borders. edges: top/bottom/left/right -> dict(sz, val, color)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "bottom", "left", "right"):
        if edge in edges:
            spec = edges[edge]
            el = tcBorders.find(qn(f"w:{edge}"))
            if el is None:
                el = OxmlElement(f"w:{edge}")
                tcBorders.append(el)
            for k, v in spec.items():
                el.set(qn(f"w:{k}"), str(v))


def add_three_line_table(doc, df, title=None, note=None, font="Times New Roman", size_pt=10):
    """Append a three-line table (and optional title/note) to an existing Document.

    Use this to drop a table into a larger document (e.g. a figure report). The
    first column is left-aligned, the rest centred — the usual convention.
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.name = font
        run.font.size = Pt(size_pt)

    table = doc.add_table(rows=1, cols=df.shape[1])
    table.autofit = True

    hdr = table.rows[0].cells
    for j, col in enumerate(df.columns):
        para = hdr[j].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(str(col))
        run.bold = True
        run.font.name = font
        run.font.size = Pt(size_pt)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, val in enumerate(row):
            para = cells[j].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run("" if val is None else str(val))
            run.font.name = font
            run.font.size = Pt(size_pt)

    thick = {"sz": 12, "val": "single", "color": "000000"}  # ~1.5 pt
    thin = {"sz": 6, "val": "single", "color": "000000"}     # ~0.75 pt
    nil = {"sz": 0, "val": "nil"}
    for cell in table.rows[0].cells:
        _set_cell_border(cell, top=thick, bottom=thin, left=nil, right=nil)
    for cell in table.rows[-1].cells:
        _set_cell_border(cell, bottom=thick, left=nil, right=nil)

    if note:
        p = doc.add_paragraph()
        run = p.add_run(note)
        run.italic = True
        run.font.name = font
        run.font.size = Pt(size_pt - 1)
    return table


def three_line_table(df, path, title=None, note=None, font="Times New Roman", size_pt=10):
    """Write `df` as a standalone three-line Word table at `path`.

    title : bold line above the table (e.g. "Table 1. ...").
    note  : smaller italic footnote below (define abbreviations, tests, n, source).
    """
    from docx import Document

    doc = Document()
    add_three_line_table(doc, df, title=title, note=note, font=font, size_pt=size_pt)
    doc.save(str(path))
    return path
